import json
import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument

from db.session import SessionLocal
from models import Document, ProspectusAnalysis

LEGAL_HEADINGS = {
    "selling restrictions",
    "definitions",
    "forward-looking statements",
    "general information",
}

DEAL_KEYWORDS = {
    "nominal value",
    "offer price range",
    "offer shares",
    "offered",
    "subscription",
    "price range",
}

PERCENT_PATTERN = re.compile(r"\b\d+(?:\.\d+)?%\b")
AED_PATTERN = re.compile(r"\bAED\s+\d+(?:,\d{3})*(?:\.\d+)?\b", re.IGNORECASE)
COMMA_INT_PATTERN = re.compile(r"\b\d{1,3}(?:,\d{3})+\b")
DATE_PATTERN = re.compile(r"\b\d{1,2}\s+[A-Za-z]+\s+\d{4}\b")


def _normalize_space(text: str) -> str:
    return " ".join(text.split())


def _guess_heading_level(paragraph) -> int | None:
    style_name = (paragraph.style.name if paragraph.style else "") or ""
    lowered = style_name.lower()
    if "heading" not in lowered:
        return None
    match = re.search(r"(\d+)", lowered)
    if not match:
        return 1
    return int(match.group(1))


def _iter_blocks(document: DocxDocument) -> Iterable[dict[str, Any]]:
    block_index = 0
    for p_index, paragraph in enumerate(document.paragraphs):
        text = _normalize_space(paragraph.text)
        if not text:
            continue
        yield {
            "block_id": f"p-{block_index}",
            "block_type": "paragraph",
            "text": text,
            "heading_level_guess": _guess_heading_level(paragraph),
            "location_path": f"document/paragraphs/{p_index}",
        }
        block_index += 1

    for t_index, table in enumerate(document.tables):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                text = _normalize_space("\n".join(p.text for p in cell.paragraphs))
                if not text:
                    continue
                yield {
                    "block_id": f"c-{block_index}",
                    "block_type": "table_cell",
                    "text": text,
                    "heading_level_guess": None,
                    "location_path": f"document/tables/{t_index}/rows/{r_index}/cells/{c_index}",
                }
                block_index += 1


def _classify_block(text: str, issuer_name: str | None, offer_shares: int | None) -> tuple[str, dict[str, Any]]:
    lowered = text.lower()
    deal_hits: list[str] = []
    boilerplate_hits: list[str] = []

    if issuer_name and issuer_name in text:
        deal_hits.append("issuer_name_exact")
    elif issuer_name and issuer_name.lower() in lowered:
        deal_hits.append("issuer_name_casefold")

    if offer_shares:
        formatted = f"{offer_shares:,}"
        if formatted in text:
            deal_hits.append("offer_shares")

    if PERCENT_PATTERN.search(text):
        deal_hits.append("percentage")
    if AED_PATTERN.search(text):
        deal_hits.append("aed_amount")
    if COMMA_INT_PATTERN.search(text):
        deal_hits.append("comma_integer")
    if DATE_PATTERN.search(text):
        deal_hits.append("date")

    if any(keyword in lowered for keyword in DEAL_KEYWORDS):
        deal_hits.append("deal_keyword")

    if any(h in lowered for h in LEGAL_HEADINGS):
        boilerplate_hits.append("legal_heading")

    tokens = re.findall(r"\w+", text)
    numeric_tokens = re.findall(r"\b\d+(?:[.,]\d+)?%?\b", text)
    numeric_density = (len(numeric_tokens) / len(tokens)) if tokens else 0.0
    if len(tokens) > 30 and numeric_density < 0.05:
        boilerplate_hits.append("dense_legal_low_numeric")

    if deal_hits and boilerplate_hits:
        classification = "mixed"
    elif deal_hits:
        classification = "deal_specific"
    else:
        classification = "boilerplate"

    return classification, {
        "deal_indicators": sorted(set(deal_hits)),
        "boilerplate_indicators": sorted(set(boilerplate_hits)),
        "numeric_density": round(numeric_density, 4),
    }


def analyze_prospectus(
    source_docx_path: str,
    issuer_name: str | None = None,
    offer_shares: int | None = None,
) -> dict[str, Any]:
    path = Path(source_docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Source DOCX not found: {source_docx_path}")

    document = DocxDocument(str(path))
    blocks: list[dict[str, Any]] = []
    counts = {"boilerplate": 0, "deal_specific": 0, "mixed": 0}

    for block in _iter_blocks(document):
        classification, signals = _classify_block(block["text"], issuer_name, offer_shares)
        block["classification"] = classification
        block["signals"] = signals
        counts[classification] += 1
        blocks.append(block)

    summary = (
        "Prospectus analysis completed: "
        f"{counts['boilerplate']} boilerplate blocks, "
        f"{counts['deal_specific']} deal-specific blocks, "
        f"{counts['mixed']} mixed blocks."
    )

    return {
        "source_docx_path": str(path),
        "total_blocks": len(blocks),
        "counts": counts,
        "summary": summary,
        "blocks": blocks,
    }


def save_analysis(
    project_id: int,
    source_document_id: int,
    analysis: dict[str, Any],
) -> int:
    session = SessionLocal()
    try:
        source_document = session.get(Document, source_document_id)
        if source_document is None:
            raise ValueError(f"Source document not found: {source_document_id}")

        row = ProspectusAnalysis(
            project_id=project_id,
            source_document_id=source_document_id,
            analysis_json=json.dumps(analysis),
        )
        session.add(row)
        session.commit()
        session.refresh(row)
        return row.id
    finally:
        session.close()
