import re
from collections.abc import Mapping

from docx.document import Document as DocxDocumentType
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([a-zA-Z0-9_.]+)\s*\}\}")
MISSING_MARKER_PATTERN = re.compile(r"\[\[MISSING:\s*([^\]]+)\]\]")


def _resolve_path(data: Mapping[str, object], path: str) -> str:
    current: object = data
    for part in path.split("."):
        if not isinstance(current, Mapping) or part not in current:
            return f"[[MISSING: {path}]]"
        current = current[part]

    if current is None:
        return f"[[MISSING: {path}]]"

    text = str(current).strip()
    return text if text else f"[[MISSING: {path}]]"


def _replace_in_runs(paragraph: Paragraph, inputs: Mapping[str, object], missing_fields: set[str]) -> None:
    if not paragraph.runs:
        return

    while True:
        runs = paragraph.runs
        run_ranges: list[tuple[int, int]] = []
        cursor = 0
        for run in runs:
            run_text = run.text
            run_ranges.append((cursor, cursor + len(run_text)))
            cursor += len(run_text)

        full_text = "".join(run.text for run in runs)
        match = PLACEHOLDER_PATTERN.search(full_text)
        if not match:
            return

        field_path = match.group(1)
        replacement = _resolve_path(inputs, field_path)
        if replacement.startswith("[[MISSING:"):
            missing_fields.add(field_path)

        span_start, span_end = match.span()
        overlap_indexes = [
            i
            for i, (start, end) in enumerate(run_ranges)
            if start < span_end and end > span_start
        ]

        if not overlap_indexes:
            return

        first_index = overlap_indexes[0]
        last_index = overlap_indexes[-1]
        first_start, _ = run_ranges[first_index]
        last_start, _ = run_ranges[last_index]

        prefix = runs[first_index].text[: max(0, span_start - first_start)]
        suffix = runs[last_index].text[max(0, span_end - last_start) :]

        runs[first_index].text = f"{prefix}{replacement}"
        for index in overlap_indexes[1:-1]:
            runs[index].text = ""
        if last_index != first_index:
            runs[last_index].text = suffix
        else:
            runs[first_index].text = f"{prefix}{replacement}{suffix}"


def _replace_in_table(table: Table, inputs: Mapping[str, object], missing_fields: set[str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            _replace_in_container(cell, inputs, missing_fields)


def _replace_in_container(
    container: DocxDocumentType | _Cell,
    inputs: Mapping[str, object],
    missing_fields: set[str],
) -> None:
    for paragraph in container.paragraphs:
        _replace_in_runs(paragraph, inputs, missing_fields)

    for table in container.tables:
        _replace_in_table(table, inputs, missing_fields)


def replace_placeholders_in_docx(document: DocxDocumentType, inputs: Mapping[str, object]) -> list[str]:
    missing_fields: set[str] = set()
    _replace_in_container(document, inputs, missing_fields)
    return sorted(missing_fields)


def extract_missing_markers(document: DocxDocumentType) -> list[str]:
    text_chunks: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_chunks.extend(paragraph.text for paragraph in cell.paragraphs)

    text = "\n".join(text_chunks)
    found = {match.group(1).strip() for match in MISSING_MARKER_PATTERN.finditer(text)}
    return sorted(found)


def extract_placeholders_from_docx(document: DocxDocumentType) -> list[str]:
    text_chunks: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_chunks.extend(paragraph.text for paragraph in cell.paragraphs)

    text = "\n".join(text_chunks)
    found = {match.group(1).strip() for match in PLACEHOLDER_PATTERN.finditer(text)}
    return sorted(found)
