import json
import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from db.session import SessionLocal
from models import Template
from services.file_service import ensure_dir, sha256_bytes
from services.prospectus_analysis_service import analyze_prospectus

TARGET_FIELDS = {
    "issuer.name": "{{issuer.name}}",
    "issuer.short_name": "{{issuer.short_name}}",
    "offer.offer_shares": "{{offer.offer_shares}}",
    "offer.percentage_offered": "{{offer.percentage_offered}}",
    "offer.price_range": "{{offer.price_range}}",
    "offer.price_range_low": "{{offer.price_range_low}}",
    "offer.price_range_high": "{{offer.price_range_high}}",
    "offer.nominal_value_per_share": "{{offer.nominal_value_per_share}}",
}


class ReplacementRule:
    def __init__(
        self,
        field: str,
        placeholder: str,
        patterns: list[re.Pattern[str]],
        requires_context: re.Pattern[str] | None = None,
    ) -> None:
        self.field = field
        self.placeholder = placeholder
        self.patterns = patterns
        self.requires_context = requires_context


NOMINAL_CONTEXT_PATTERN = re.compile(r"nominal\s+value(?:\s+per\s+share)?", re.IGNORECASE)
SHORT_NAME_INFERENCE_PATTERN = re.compile(
    r"\(\s*the\s+[\"'‘’]Company[\"'‘’]\s+or\s+[\"'‘’](?P<short>[A-Za-z0-9][^\"'‘’]{1,50})[\"'‘’]\s*\)",
    re.IGNORECASE,
)


def _iter_text_blocks(document: DocxDocument) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for container, container_path in _iter_containers(document):
        for p_index, paragraph in enumerate(container.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            if container_path == "document":
                location_path = f"document/paragraphs/{p_index}"
            else:
                location_path = f"{container_path}/paragraphs/{p_index}"
            blocks.append((location_path, text))
    return blocks


def extract_source_deal_values(source_docx_path: str) -> dict[str, Any]:
    document = DocxDocument(source_docx_path)
    blocks = _iter_text_blocks(document)

    patterns: dict[str, tuple[re.Pattern[str], Any]] = {
        "issuer.name": (re.compile(r"\b([A-Z][A-Za-z0-9&\-. ]+?\s+(?:plc|PJSC|LLC|L\.L\.C\.))\b"), str),
        "offer.offer_shares": (re.compile(r"offer\s+shares\s*:\s*([\d,]+)", re.IGNORECASE), int),
        "offer.percentage_offered": (
            re.compile(r"percentage\s+offered\s*:\s*([\d.]+)%", re.IGNORECASE),
            float,
        ),
        "offer.nominal_value_per_share_aed": (
            re.compile(r"nominal\s+value\s+per\s+share\s*:\s*AED\s*([\d.]+)", re.IGNORECASE),
            float,
        ),
        "offer.price_range_low_aed": (
            re.compile(r"offer\s+price\s+range\s*:\s*AED\s*([\d.]+)\s*[\-–—]\s*AED\s*([\d.]+)", re.IGNORECASE),
            float,
        ),
        "offer.price_range_high_aed": (
            re.compile(r"offer\s+price\s+range\s*:\s*AED\s*([\d.]+)\s*[\-–—]\s*AED\s*([\d.]+)", re.IGNORECASE),
            float,
        ),
    }

    values: dict[str, Any] = {}
    evidence: dict[str, list[dict[str, Any]]] = {}

    for field, (pattern, caster) in patterns.items():
        for location_path, text in blocks:
            match = pattern.search(text)
            if not match:
                continue
            raw = match.group(1)
            if field == "offer.price_range_high_aed":
                raw = match.group(2)
            parsed: Any = raw
            if caster is int:
                parsed = int(str(raw).replace(",", ""))
            elif caster is float:
                parsed = float(raw)
            elif caster is str:
                parsed = str(raw).strip()

            values[field] = parsed
            evidence.setdefault(field, []).append(
                {
                    "location_path": location_path,
                    "snippet": text[:180],
                    "confidence": 0.92,
                }
            )
            break

    return {"values": values, "evidence": evidence}


def _nested_payload(flat_values: dict[str, Any]) -> dict[str, Any]:
    payload: dict[str, Any] = {}
    for key, value in flat_values.items():
        current = payload
        parts = key.split(".")
        for part in parts[:-1]:
            current = current.setdefault(part, {})
        current[parts[-1]] = value
    return payload


def _merge_dicts(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    result = dict(base)
    for key, value in override.items():
        if key in result and isinstance(result[key], dict) and isinstance(value, dict):
            result[key] = _merge_dicts(result[key], value)
        elif value not in (None, "", {}):
            result[key] = value
    return result


def _replace_match_in_runs(paragraph: Paragraph, pattern: re.Pattern[str], replacement: str) -> int:
    if not paragraph.runs:
        return 0

    replacements = 0
    while True:
        runs = paragraph.runs
        run_ranges: list[tuple[int, int]] = []
        cursor = 0
        for run in runs:
            run_text = run.text
            run_ranges.append((cursor, cursor + len(run_text)))
            cursor += len(run_text)

        full_text = "".join(run.text for run in runs)
        match = pattern.search(full_text)
        if not match:
            break

        span_start, span_end = match.span()
        overlap_indexes = [
            i
            for i, (start, end) in enumerate(run_ranges)
            if start < span_end and end > span_start
        ]
        if not overlap_indexes:
            break

        first_index = overlap_indexes[0]
        last_index = overlap_indexes[-1]
        first_start, _ = run_ranges[first_index]
        last_start, _ = run_ranges[last_index]

        prefix = runs[first_index].text[: max(0, span_start - first_start)]
        suffix = runs[last_index].text[max(0, span_end - last_start) :]

        runs[first_index].text = f"{prefix}{replacement}"
        for index in overlap_indexes[1:-1]:
            runs[index].text = ""
        if last_index != first_index:
            runs[last_index].text = suffix
        else:
            runs[first_index].text = f"{prefix}{replacement}{suffix}"

        replacements += 1

    return replacements


def _iter_containers(document: DocxDocument):
    yield document, "document"
    for t_index, table in enumerate(document.tables):
        yield from _iter_table(table, f"document/tables/{t_index}")
    for section_index, section in enumerate(document.sections):
        section_roots = [
            (section.header, f"sections/{section_index}/header/default"),
            (section.first_page_header, f"sections/{section_index}/header/first_page"),
            (section.even_page_header, f"sections/{section_index}/header/even_page"),
            (section.footer, f"sections/{section_index}/footer/default"),
            (section.first_page_footer, f"sections/{section_index}/footer/first_page"),
            (section.even_page_footer, f"sections/{section_index}/footer/even_page"),
        ]
        for root, path in section_roots:
            yield root, path
            for t_index, table in enumerate(root.tables):
                yield from _iter_table(table, f"{path}/tables/{t_index}")


def _iter_table(table: Table, table_path: str):
    for r_index, row in enumerate(table.rows):
        for c_index, cell in enumerate(row.cells):
            cell_path = f"{table_path}/rows/{r_index}/cells/{c_index}"
            yield cell, cell_path
            for nested_t_index, nested in enumerate(cell.tables):
                yield from _iter_table(nested, f"{cell_path}/tables/{nested_t_index}")


def _count_pattern_matches(pattern: re.Pattern[str], text: str) -> int:
    return sum(1 for _ in pattern.finditer(text))


def _apply_rule(
    document: DocxDocument,
    rule: ReplacementRule,
    allowed_paths: set[str],
) -> dict[str, Any]:
    report = {
        "field": rule.field,
        "placeholder": rule.placeholder,
        "found_count": 0,
        "replaced_count": 0,
        "skipped_count": 0,
        "sample_locations": [],
    }
    block_counter = 0
    for container, container_path in _iter_containers(document):
        for p_index, paragraph in enumerate(container.paragraphs):
            if container_path == "document":
                location_path = f"document/paragraphs/{p_index}"
            else:
                location_path = f"{container_path}/paragraphs/{p_index}"

            if rule.requires_context and not rule.requires_context.search(paragraph.text):
                block_counter += 1
                continue

            found_here = sum(_count_pattern_matches(pattern, paragraph.text) for pattern in rule.patterns)
            if found_here == 0:
                block_counter += 1
                continue

            report["found_count"] += found_here
            if len(report["sample_locations"]) < 5:
                report["sample_locations"].append(
                    {"block_id": f"block-{block_counter}", "location_path": location_path}
                )

            if allowed_paths:
                is_allowed = any(
                    location_path == allowed or location_path.startswith(f"{allowed}/")
                    for allowed in allowed_paths
                )
                if not is_allowed:
                    report["skipped_count"] += found_here
                    block_counter += 1
                    continue

            total_here = 0
            for pattern in rule.patterns:
                total_here += _replace_match_in_runs(paragraph, pattern, rule.placeholder)
            report["replaced_count"] += total_here
            block_counter += 1
    return report


def _normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _number_variants(value: float) -> list[str]:
    rounded = round(float(value), 3)
    base = f"{rounded:.3f}".rstrip("0").rstrip(".")
    variants = {
        base,
        f"{rounded:.1f}",
        f"{rounded:.2f}",
        f"{rounded:.3f}",
    }
    return sorted({v for v in variants if re.match(r"^\d+(?:\.\d{1,3})?$", v)})


def _build_number_pattern(value: float) -> str:
    return "(?:" + "|".join(re.escape(v) for v in _number_variants(value)) + ")"


def _infer_issuer_short_name(document: DocxDocument) -> str | None:
    for container, _ in _iter_containers(document):
        for paragraph in container.paragraphs:
            match = SHORT_NAME_INFERENCE_PATTERN.search(paragraph.text)
            if match:
                short_name = _normalize_whitespace(match.group("short"))
                if short_name and short_name.lower() != "company":
                    return short_name
    return None


def _build_rules(document: DocxDocument, inputs: dict[str, Any]) -> list[ReplacementRule]:
    issuer_name = str(inputs.get("issuer", {}).get("name") or "").strip()
    issuer_short_name = str(inputs.get("issuer", {}).get("short_name") or "").strip()
    offer = inputs.get("offer", {})

    rules: list[ReplacementRule] = []

    if issuer_name:
        normalized_issuer = _normalize_whitespace(issuer_name)
        issuer_pattern = re.sub(r"\s+", r"\\s+", re.escape(normalized_issuer))
        rules.append(
            ReplacementRule(
                field="issuer.name",
                placeholder=TARGET_FIELDS["issuer.name"],
                patterns=[
                    re.compile(re.escape(issuer_name)),
                    re.compile(issuer_pattern, re.IGNORECASE),
                ],
            )
        )

    inferred_short_name = _infer_issuer_short_name(document)
    short_name = issuer_short_name or inferred_short_name
    if short_name:
        rules.append(
            ReplacementRule(
                field="issuer.short_name",
                placeholder=TARGET_FIELDS["issuer.short_name"],
                patterns=[re.compile(rf"\b{re.escape(short_name)}\b", re.IGNORECASE)],
            )
        )

    offer_shares = offer.get("offer_shares")
    if offer_shares is not None:
        formatted = f"{int(offer_shares):,}"
        rules.append(
            ReplacementRule(
                field="offer.offer_shares",
                placeholder=TARGET_FIELDS["offer.offer_shares"],
                patterns=[re.compile(re.escape(formatted))],
            )
        )

    percentage = offer.get("percentage_offered")
    if percentage is not None:
        percent_text = f"{float(percentage):g}%"
        rules.append(
            ReplacementRule(
                field="offer.percentage_offered",
                placeholder=TARGET_FIELDS["offer.percentage_offered"],
                patterns=[re.compile(re.escape(percent_text))],
            )
        )

    nominal = offer.get("nominal_value_per_share_aed")
    if nominal is not None:
        nominal_number_pattern = _build_number_pattern(float(nominal))
        rules.append(
            ReplacementRule(
                field="offer.nominal_value_per_share",
                placeholder=TARGET_FIELDS["offer.nominal_value_per_share"],
                patterns=[re.compile(rf"\bAED\s*{nominal_number_pattern}\b", re.IGNORECASE)],
                requires_context=NOMINAL_CONTEXT_PATTERN,
            )
        )

    low = offer.get("price_range_low_aed")
    high = offer.get("price_range_high_aed")
    if low is not None and high is not None:
        low_pattern = _build_number_pattern(float(low))
        high_pattern = _build_number_pattern(float(high))
        rules.append(
            ReplacementRule(
                field="offer.price_range",
                placeholder=TARGET_FIELDS["offer.price_range"],
                patterns=[
                    re.compile(
                        rf"\bAED\s*{low_pattern}\s*(?:[\-–—]\s*(?:AED\s*)?|to\s+(?:AED\s*)?){high_pattern}\b",
                        re.IGNORECASE,
                    )
                ],
            )
        )
        rules.append(
            ReplacementRule(
                field="offer.price_range_low",
                placeholder=TARGET_FIELDS["offer.price_range_low"],
                patterns=[re.compile(rf"\bAED\s*{low_pattern}\b", re.IGNORECASE)],
            )
        )
        rules.append(
            ReplacementRule(
                field="offer.price_range_high",
                placeholder=TARGET_FIELDS["offer.price_range_high"],
                patterns=[re.compile(rf"\bAED\s*{high_pattern}\b", re.IGNORECASE)],
            )
        )

    return rules


def _count_potentially_skipped_textboxes(document: DocxDocument) -> int:
    parts = [document.part]
    for section in document.sections:
        parts.extend(
            [
                section.header.part,
                section.first_page_header.part,
                section.even_page_header.part,
                section.footer.part,
                section.first_page_footer.part,
                section.even_page_footer.part,
            ]
        )

    unique_parts = {part.partname: part for part in parts}.values()
    return sum(len(part._element.xpath(".//*[local-name()='txbxContent']")) for part in unique_parts)


def _next_template_version(session, template_name: str) -> int:
    latest = (
        session.query(Template)
        .filter(Template.name == template_name)
        .order_by(Template.version.desc())
        .first()
    )
    return 1 if latest is None else latest.version + 1


def _safe_stem(name: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "_", name.strip())
    return cleaned.strip("_") or "parameterized_template"


def parameterize_template_from_source(
    source_docx_path: str,
    inputs: dict[str, Any],
    base_template_id: int,
    source_document_id: int,
    project_id: int,
    aliases: dict[str, list[str]] | None = None,
    dry_run: bool = False,
) -> dict[str, Any]:
    extracted = extract_source_deal_values(source_docx_path)
    merged_inputs = _merge_dicts(_nested_payload(extracted["values"]), inputs)

    _ = aliases
    analysis = analyze_prospectus(
        source_docx_path,
        issuer_name=str(merged_inputs.get("issuer", {}).get("name") or "").strip() or None,
        offer_shares=merged_inputs.get("offer", {}).get("offer_shares"),
    )
    allowed_paths = {
        block["location_path"]
        for block in analysis["blocks"]
        if block["classification"] in {"deal_specific", "mixed"}
    }

    document = DocxDocument(source_docx_path)
    rules = _build_rules(document, merged_inputs)
    field_reports = {field: {"found_count": 0, "replaced_count": 0, "skipped_count": 0, "sample_locations": []} for field in TARGET_FIELDS}

    for rule in rules:
        rule_report = _apply_rule(document, rule, allowed_paths)
        field_reports[rule.field] = {
            "found_count": rule_report["found_count"],
            "replaced_count": rule_report["replaced_count"],
            "skipped_count": rule_report["skipped_count"],
            "sample_locations": rule_report["sample_locations"],
        }

    placeholder_count = sum(report["replaced_count"] for report in field_reports.values())
    placeholders_replaced = sorted(
        TARGET_FIELDS[field]
        for field, report in field_reports.items()
        if report["replaced_count"] > 0
    )
    unresolved = sorted(
        TARGET_FIELDS[field] for field, report in field_reports.items() if report["found_count"] == 0
    )
    potentially_skipped_textboxes = _count_potentially_skipped_textboxes(document)
    notes: list[str] = []
    if potentially_skipped_textboxes > 0:
        notes.append(
            "python-docx cannot reliably edit text boxes/shapes; replacements there may be skipped."
        )

    if placeholder_count == 0:
        raise ValueError(
            "No placeholders were inserted. This is likely a static prospectus with unmatched patterns for the provided inputs."
        )

    report = {
        "placeholder_count": placeholder_count,
        "placeholders": placeholders_replaced,
        "fields": field_reports,
        "unresolved": unresolved,
        "potentially_skipped": {"textboxes_or_shapes": potentially_skipped_textboxes},
        "notes": notes,
    }

    if dry_run:
        return {
            "template_id": None,
            "new_template_docx_path": None,
            "parameterization_report": report,
            "analysis": analysis,
            "source_extraction": extracted,
        }

    session = SessionLocal()
    try:
        base_template = session.get(Template, base_template_id)
        if base_template is None:
            raise ValueError(f"Base template not found: {base_template_id}")

        parameterized_name = f"{base_template.name} (parameterized)"
        next_version = _next_template_version(session, parameterized_name)
        output_dir = ensure_dir(Path("storage") / "templates" / "parameterized")
        output_name = f"{_safe_stem(base_template.name)}_v{next_version}.docx"
        output_path = output_dir / output_name

        document.save(str(output_path))
        file_bytes = output_path.read_bytes()

        metadata = {
            "source_template_id": base_template_id,
            "parameterized_from_document_id": source_document_id,
            "placeholder_count": placeholder_count,
            "analysis_counts": analysis["counts"],
            "source_extraction": extracted,
            "parameterization": report,
        }

        template = Template(
            name=parameterized_name,
            status="draft",
            sha256=sha256_bytes(file_bytes),
            file_path=str(output_path),
            version=next_version,
            metadata_json=json.dumps(metadata),
        )
        session.add(template)
        session.commit()
        session.refresh(template)

        return {
            "template_id": template.id,
            "new_template_docx_path": str(output_path),
            "parameterization_report": report,
            "analysis": analysis,
            "source_extraction": extracted,
        }
    finally:
        session.close()
