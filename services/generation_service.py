import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from db.session import SessionLocal
from models import Document, GenerationRun, Template
from services.file_service import ensure_dir, sha256_bytes
from services.normalization_service import normalize_inputs
from services.placeholder_service import extract_missing_markers, replace_placeholders_in_docx


def _insert_paragraph_at_start(document: DocxDocument, text: str) -> None:
    if not document.paragraphs:
        document.add_paragraph(text)
        return

    first_paragraph = document.paragraphs[0]
    paragraph_xml = OxmlElement("w:p")
    first_paragraph._p.addprevious(paragraph_xml)  # noqa: SLF001
    paragraph = Paragraph(paragraph_xml, first_paragraph._parent)  # noqa: SLF001
    paragraph.add_run(text)


def _prepend_missing_information(document: DocxDocument, missing_fields: list[str]) -> None:
    if not missing_fields:
        return

    lines = [f"- [[MISSING: {field}]]" for field in missing_fields]
    lines.insert(0, "Missing Information")
    lines.insert(1, "")

    for line in reversed(lines):
        _insert_paragraph_at_start(document, line)


def _next_document_version(session, project_id: int, doc_type: str) -> int:
    latest = (
        session.query(Document)
        .filter(Document.project_id == project_id, Document.doc_type == doc_type)
        .order_by(Document.version.desc())
        .first()
    )
    return 1 if latest is None else latest.version + 1


def _create_source_document_from_template(
    session,
    project_id: int,
    template_id: int,
    template_path: Path,
) -> int:
    source_dir = ensure_dir(Path("storage") / "projects" / str(project_id) / "source")
    source_version = _next_document_version(session, project_id, "original")
    source_name = f"source_v{source_version}_template_{template_id}.docx"
    source_path = source_dir / source_name
    shutil.copyfile(template_path, source_path)

    source_sha256 = sha256_bytes(source_path.read_bytes())
    source_document = Document(
        project_id=project_id,
        doc_type="original",
        file_name=source_name,
        file_path=str(source_path),
        sha256=source_sha256,
        version=source_version,
        is_locked=False,
    )
    session.add(source_document)
    session.flush()
    return source_document.id


def generate_draft_docx(project_id: int, template_id: int, inputs_json: str | dict[str, Any]) -> dict[str, Any]:
    inputs_payload = inputs_json if isinstance(inputs_json, dict) else json.loads(inputs_json)
    schema_id = str(inputs_payload.get("schema_id") or "talabat_v1")
    normalized_inputs, rendered_fields_map, normalization_missing = normalize_inputs(schema_id, inputs_payload)

    session = SessionLocal()
    try:
        template = session.get(Template, template_id)
        if template is None:
            raise ValueError(f"Template not found: {template_id}")

        template_path = Path(template.file_path)
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template.file_path}")

        source_document_id = normalized_inputs.get("source_document_id")
        if source_document_id is None and normalized_inputs.get("use_template_as_source"):
            source_document_id = _create_source_document_from_template(
                session=session,
                project_id=project_id,
                template_id=template_id,
                template_path=template_path,
            )
            normalized_inputs["source_document_id"] = source_document_id

        document = DocxDocument(str(template_path))
        replaced_missing = replace_placeholders_in_docx(document, normalized_inputs)
        marker_missing = extract_missing_markers(document)
        all_missing_fields = sorted(set(replaced_missing + marker_missing + normalization_missing))
        _prepend_missing_information(document, all_missing_fields)

        output_dir = ensure_dir(Path("storage") / "projects" / str(project_id) / "generated")
        next_version = _next_document_version(session, project_id, "draft")
        output_name = f"draft_v{next_version}_template_{template_id}.docx"
        output_path = output_dir / output_name
        document.save(str(output_path))

        output_sha256 = sha256_bytes(output_path.read_bytes())

        draft_document = Document(
            project_id=project_id,
            doc_type="draft",
            file_name=output_name,
            file_path=str(output_path),
            sha256=output_sha256,
            version=next_version,
            is_locked=False,
        )
        session.add(draft_document)
        session.flush()

        run = (
            session.query(GenerationRun)
            .filter(
                GenerationRun.project_id == project_id,
                GenerationRun.template_id == template_id,
                GenerationRun.source_document_id == source_document_id,
                GenerationRun.status == "pending",
            )
            .order_by(GenerationRun.created_at.desc())
            .first()
        )

        if run is None:
            run = GenerationRun(
                project_id=project_id,
                template_id=template_id,
                source_document_id=source_document_id,
                status="pending",
                inputs_json=json.dumps(normalized_inputs),
            )
            session.add(run)
            session.flush()

        run.output_document_id = draft_document.id
        run.output_path = str(output_path)
        run.status = "completed"

        session.commit()

        return {
            "document_id": draft_document.id,
            "output_path": str(output_path),
            "missing_fields": all_missing_fields,
            "generation_run_id": run.id,
            "rendered_fields": rendered_fields_map,
            "source_document_id": source_document_id,
        }
    finally:
        session.close()
