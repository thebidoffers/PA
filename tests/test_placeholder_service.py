import pytest

pytest.importorskip("docx")

from pathlib import Path

from docx import Document as DocxDocument

from services.placeholder_service import extract_placeholders_from_docx, replace_placeholders_in_docx


def test_replace_placeholders_in_paragraphs_and_tables(tmp_path):
    template_path = Path(tmp_path) / "placeholder_fixture.docx"

    document = DocxDocument()
    paragraph = document.add_paragraph()
    paragraph.add_run("Issuer: {{issuer")
    paragraph.add_run(".name}}")
    paragraph.add_run(" | Country: {{issuer.country}}")

    table = document.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Offer size: {{offer.size}}"

    document.save(str(template_path))

    loaded = DocxDocument(str(template_path))
    missing_fields = replace_placeholders_in_docx(
        loaded,
        {
            "issuer": {"name": "Acme Holdings"},
        },
    )

    assert loaded.paragraphs[0].text == "Issuer: Acme Holdings | Country: [[MISSING: issuer.country]]"
    assert loaded.tables[0].rows[0].cells[0].text == "Offer size: [[MISSING: offer.size]]"
    assert missing_fields == ["issuer.country", "offer.size"]



def test_extract_placeholders_from_docx_reads_paragraphs_and_tables(tmp_path):
    fixture_path = Path(tmp_path) / "extract_fixture.docx"

    document = DocxDocument()
    document.add_paragraph("Issuer: {{issuer.name}}")
    table = document.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Offer size: {{offer.offer_shares}}"
    document.save(str(fixture_path))

    loaded = DocxDocument(str(fixture_path))
    placeholders = extract_placeholders_from_docx(loaded)

    assert placeholders == ["issuer.name", "offer.offer_shares"]
