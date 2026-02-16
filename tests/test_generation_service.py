import pytest

pytest.importorskip("docx")

import importlib
import sys
import json
from pathlib import Path

from docx import Document as DocxDocument


def test_generate_draft_docx_creates_file_and_db_rows(tmp_path, monkeypatch):
    db_file = tmp_path / "generation.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{db_file}")

    session_module = importlib.import_module("db.session")
    importlib.reload(session_module)

    session_module.Base.metadata.clear()
    sys.modules.pop("models", None)
    sys.modules.pop("models.entities", None)
    models_module = importlib.import_module("models.entities")

    generation_service = importlib.import_module("services.generation_service")
    importlib.reload(generation_service)

    base = session_module.Base
    engine = session_module.engine
    base.metadata.create_all(bind=engine)

    template_path = tmp_path / "template.docx"
    template_doc = DocxDocument()
    template_doc.add_paragraph("Issuer: {{issuer.name}}")
    template_doc.add_paragraph("Offer Shares: {{offer.offer_shares}}")
    template_doc.add_paragraph("Legacy Offer Size: {{offer.size}}")
    template_doc.add_paragraph("Offer Price Range: {{offer.price_range}}")
    template_doc.add_paragraph("Undisclosed: {{issuer.country}}")
    template_doc.save(str(template_path))

    session = session_module.SessionLocal()
    try:
        project = models_module.ProspectusProject(name="Project Gen")
        session.add(project)
        session.flush()

        template = models_module.Template(
            name="Draft Template",
            status="approved",
            sha256="a" * 64,
            file_path=str(template_path),
        )
        session.add(template)
        session.flush()

        source_document = models_module.Document(
            project_id=project.id,
            doc_type="original",
            file_name="source.docx",
            file_path=str(template_path),
            sha256="b" * 64,
            version=1,
        )
        session.add(source_document)
        session.flush()

        run = models_module.GenerationRun(
            project_id=project.id,
            template_id=template.id,
            source_document_id=source_document.id,
            status="pending",
            inputs_json=json.dumps({"seed": "value"}),
            output_path=None,
        )
        session.add(run)
        session.commit()

        result = generation_service.generate_draft_docx(
            project.id,
            template.id,
            {
                "schema_id": "talabat_v1",
                "issuer": {"name": "Acme Holdings"},
                "offer": {
                    "offer_shares": 3493236093,
                    "price_range_low_aed": 1.3,
                    "price_range_high_aed": 1.5,
                },
                "source_document_id": source_document.id,
            },
        )

        output_path = Path(result["output_path"])
        assert output_path.exists()

        created_document = session.get(models_module.Document, result["document_id"])
        assert created_document is not None
        assert created_document.doc_type == "draft"
        assert created_document.version == 1

        updated_run = session.get(models_module.GenerationRun, run.id)
        assert updated_run.status == "completed"
        assert updated_run.output_document_id == created_document.id

        generated_doc = DocxDocument(str(output_path))
        generated_text = "\n".join(paragraph.text for paragraph in generated_doc.paragraphs)
        assert "3,493,236,093" in generated_text
        assert "AED 1.30 – AED 1.50" in generated_text
        assert "Missing Information" in generated_text
        assert "[[MISSING: issuer.country]]" in generated_text
        assert "key_dates" not in result["missing_fields"]
        assert sorted(result["template_placeholders"]) == [
            "issuer.country",
            "issuer.name",
            "offer.offer_shares",
            "offer.price_range",
            "offer.size",
        ]
    finally:
        session.close()


def test_generate_draft_docx_supports_template_as_source(tmp_path, monkeypatch):
    db_file = tmp_path / "generation_template_source.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{db_file}")

    session_module = importlib.import_module("db.session")
    importlib.reload(session_module)

    session_module.Base.metadata.clear()
    sys.modules.pop("models", None)
    sys.modules.pop("models.entities", None)
    models_module = importlib.import_module("models.entities")

    generation_service = importlib.import_module("services.generation_service")
    importlib.reload(generation_service)

    base = session_module.Base
    engine = session_module.engine
    base.metadata.create_all(bind=engine)

    template_path = tmp_path / "template_source.docx"
    template_doc = DocxDocument()
    template_doc.add_paragraph("Issuer: {{issuer.name}}")
    template_doc.add_paragraph("Offer Shares: {{offer.offer_shares}}")
    template_doc.add_paragraph("Offer Price Range: {{offer.price_range}}")
    template_doc.save(str(template_path))

    session = session_module.SessionLocal()
    try:
        project = models_module.ProspectusProject(name="Project No Docs")
        session.add(project)
        session.flush()

        template = models_module.Template(
            name="Template Source",
            status="approved",
            sha256="c" * 64,
            file_path=str(template_path),
        )
        session.add(template)
        session.commit()

        result = generation_service.generate_draft_docx(
            project.id,
            template.id,
            {
                "schema_id": "talabat_v1",
                "issuer": {"name": "Acme Holdings"},
                "offer": {
                    "offer_shares": 1000,
                    "price_range_low_aed": 1.3,
                    "price_range_high_aed": 1.5,
                },
                "use_template_as_source": True,
            },
        )

        run = session.get(models_module.GenerationRun, result["generation_run_id"])
        assert run is not None
        assert run.status == "completed"
        assert run.source_document_id is not None

        source_document = session.get(models_module.Document, run.source_document_id)
        assert source_document is not None
        assert source_document.doc_type == "original"
        assert Path(source_document.file_path).exists()

        output_path = Path(result["output_path"])
        assert output_path.exists()
    finally:
        session.close()


def test_generation_uses_normalized_values_and_deal_profile_persists(tmp_path, monkeypatch):
    db_file = tmp_path / "generation_profile.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{db_file}")

    session_module = importlib.import_module("db.session")
    importlib.reload(session_module)

    session_module.Base.metadata.clear()
    sys.modules.pop("models", None)
    sys.modules.pop("models.entities", None)
    models_module = importlib.import_module("models.entities")

    init_db_module = importlib.import_module("db.init_db")
    importlib.reload(init_db_module)
    init_db_module.init_db()

    generation_service = importlib.import_module("services.generation_service")
    importlib.reload(generation_service)
    normalization_service = importlib.import_module("services.normalization_service")
    deal_profile_service = importlib.import_module("services.deal_profile_service")
    importlib.reload(deal_profile_service)

    template_path = tmp_path / "template_profile.docx"
    template_doc = DocxDocument()
    template_doc.add_paragraph("Offer Shares: {{offer.offer_shares}}")
    template_doc.add_paragraph("Offer Price Range: {{offer.price_range}}")
    template_doc.save(str(template_path))

    session = session_module.SessionLocal()
    try:
        project = models_module.ProspectusProject(name="Project Profile Flow")
        session.add(project)
        session.flush()

        template = models_module.Template(
            name="Profile Template",
            status="approved",
            sha256="d" * 64,
            file_path=str(template_path),
        )
        session.add(template)
        session.commit()

        raw_inputs = {
            "schema_id": "talabat_v1",
            "issuer": {"name": "Issuer X"},
            "offer": {
                "offer_shares": "3493236093",
                "price_range_low_aed": "1.30",
                "price_range_high_aed": "1.50",
            },
            "use_template_as_source": True,
        }
        normalized, _, _ = normalization_service.normalize_inputs("talabat_v1", raw_inputs)

        saved = deal_profile_service.save_profile(
            project_id=project.id,
            schema_id="talabat_v1",
            template_id=template.id,
            inputs_raw=raw_inputs,
            inputs_normalized=normalized,
        )

        result = generation_service.generate_draft_docx(project.id, template.id, normalized)

        generated_doc = DocxDocument(result["output_path"])
        generated_text = "\n".join(paragraph.text for paragraph in generated_doc.paragraphs)

        assert "3,493,236,093" in generated_text
        assert "AED 1.30 – AED 1.50" in generated_text

        latest_profile = deal_profile_service.get_latest_profile(project.id, "talabat_v1", template.id)
        assert latest_profile is not None
        assert latest_profile.id == saved.id
        persisted = json.loads(latest_profile.inputs_normalized_json)
        assert persisted["offer"]["price_range"] == "AED 1.30 – AED 1.50"
    finally:
        session.close()
