import importlib
import json
import sys

import pytest

pytest.importorskip("docx")

from docx import Document as DocxDocument

from tests.utils.docx_factory import make_talabat_like_docx


def test_parameterize_template_from_source_replaces_targeted_fields_with_coverage(tmp_path, monkeypatch):
    db_file = tmp_path / "parameterization.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{db_file}")

    session_module = importlib.import_module("db.session")
    importlib.reload(session_module)

    session_module.Base.metadata.clear()
    sys.modules.pop("models", None)
    sys.modules.pop("models.entities", None)
    models_module = importlib.import_module("models.entities")

    parameterization_service = importlib.import_module("services.parameterization_service")
    importlib.reload(parameterization_service)

    base = session_module.Base
    engine = session_module.engine
    base.metadata.create_all(bind=engine)

    source_path = make_talabat_like_docx(tmp_path / "source.docx")

    session = session_module.SessionLocal()
    try:
        project = models_module.ProspectusProject(name="Parameterization Project")
        session.add(project)
        session.flush()

        source_doc = models_module.Document(
            project_id=project.id,
            doc_type="original",
            file_name="source.docx",
            file_path=str(source_path),
            sha256="a" * 64,
            version=1,
        )
        session.add(source_doc)
        session.flush()

        base_template = models_module.Template(
            name="Talabat Base",
            status="approved",
            sha256="b" * 64,
            file_path=str(source_path),
            version=1,
        )
        session.add(base_template)
        session.commit()

        result = parameterization_service.parameterize_template_from_source(
            source_docx_path=str(source_path),
            inputs={
                "issuer": {"name": "Talabat Holding plc"},
                "offer": {
                    "offer_shares": 3493236093,
                    "percentage_offered": 15,
                    "nominal_value_per_share_aed": 1.0,
                    "price_range_low_aed": 1.30,
                    "price_range_high_aed": 1.50,
                },
            },
            base_template_id=base_template.id,
            source_document_id=source_doc.id,
            project_id=project.id,
        )

        report = result["parameterization_report"]
        assert report["placeholder_count"] > 2
        assert "{{issuer.name}}" in report["placeholders"]
        assert "{{offer.price_range}}" in report["placeholders"]
        assert report["fields"]["issuer.name"]["replaced_count"] > 0
        assert report["fields"]["offer.price_range"]["replaced_count"] > 0
        assert report["fields"]["offer.nominal_value_per_share"]["replaced_count"] > 0
        extraction = result["source_extraction"]
        assert extraction["values"]["issuer.name"]
        assert extraction["values"]["offer.nominal_value_per_share_aed"] == 1.0
        for field in ("issuer.name", "offer.nominal_value_per_share_aed"):
            assert extraction["evidence"][field]
            assert extraction["evidence"][field][0]["snippet"]
            assert extraction["evidence"][field][0]["confidence"] > 0

        generated = DocxDocument(result["new_template_docx_path"])
        output_text = "\n".join(p.text for p in generated.paragraphs)
        output_header_text = "\n".join(p.text for p in generated.sections[0].header.paragraphs)

        assert "{{issuer.name}}" in output_text or "{{issuer.name}}" in output_header_text
        assert "{{issuer.short_name}}" in output_text
        assert "{{offer.offer_shares}}" in output_text
        assert "{{offer.percentage_offered}}" in output_text
        assert "{{offer.nominal_value_per_share}}" in output_text
        assert "{{offer.price_range}}" in output_text
        assert "{{offer.price_range_low}}" in output_text
        assert "{{offer.price_range_high}}" in output_text

        created_template = session.get(models_module.Template, result["template_id"])
        assert created_template is not None
        metadata = json.loads(created_template.metadata_json)
        assert metadata["placeholder_count"] == report["placeholder_count"]
        assert metadata["source_template_id"] == base_template.id
        assert "source_extraction" in metadata
        assert "parameterization" in metadata
        assert metadata["parameterization"]["fields"]["issuer.name"]["replaced_count"] > 0
    finally:
        session.close()


def test_parameterize_template_from_source_dry_run_returns_report_without_writing(tmp_path, monkeypatch):
    db_file = tmp_path / "parameterization_dry_run.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{db_file}")

    session_module = importlib.import_module("db.session")
    importlib.reload(session_module)

    session_module.Base.metadata.clear()
    sys.modules.pop("models", None)
    sys.modules.pop("models.entities", None)
    models_module = importlib.import_module("models.entities")

    parameterization_service = importlib.import_module("services.parameterization_service")
    importlib.reload(parameterization_service)

    base = session_module.Base
    engine = session_module.engine
    base.metadata.create_all(bind=engine)

    source_path = make_talabat_like_docx(tmp_path / "source_dry_run.docx")

    session = session_module.SessionLocal()
    try:
        project = models_module.ProspectusProject(name="Dry Run Project")
        session.add(project)
        session.flush()

        source_doc = models_module.Document(
            project_id=project.id,
            doc_type="original",
            file_name="source_dry_run.docx",
            file_path=str(source_path),
            sha256="c" * 64,
            version=1,
        )
        session.add(source_doc)
        session.flush()

        base_template = models_module.Template(
            name="Dry Run Base",
            status="approved",
            sha256="d" * 64,
            file_path=str(source_path),
            version=1,
        )
        session.add(base_template)
        session.commit()

        result = parameterization_service.parameterize_template_from_source(
            source_docx_path=str(source_path),
            inputs={"issuer": {}, "offer": {}},
            base_template_id=base_template.id,
            source_document_id=source_doc.id,
            project_id=project.id,
            dry_run=True,
        )

        assert result["template_id"] is None
        assert result["new_template_docx_path"] is None
        assert result["parameterization_report"]["placeholder_count"] > 0
        assert result["source_extraction"]["values"]["issuer.name"]
        assert result["source_extraction"]["values"]["offer.nominal_value_per_share_aed"] == 1.0
        assert session.query(models_module.Template).count() == 1
    finally:
        session.close()


def test_extract_source_deal_values_returns_evidence_and_confidence(tmp_path):
    parameterization_service = importlib.import_module("services.parameterization_service")
    importlib.reload(parameterization_service)

    source_path = make_talabat_like_docx(tmp_path / "source_extraction.docx")
    extraction = parameterization_service.extract_source_deal_values(str(source_path))

    assert extraction["values"]["issuer.name"]
    assert extraction["values"]["offer.nominal_value_per_share_aed"] == 1.0
    for field in ("issuer.name", "offer.nominal_value_per_share_aed"):
        evidence_items = extraction["evidence"][field]
        assert len(evidence_items) >= 1
        assert evidence_items[0]["snippet"]
        assert evidence_items[0]["confidence"] > 0
