from pathlib import Path

from docx import Document as DocxDocument


def make_talabat_like_docx(path: str | Path) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)

    document = DocxDocument()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Header token issuer: Talabat Holding plc"
    section.footer.paragraphs[0].text = "Footer token nominal value: AED 1.00"

    document.add_paragraph("Talabat Holding plc (the 'Company' or 'talabat') is offering shares.")
    document.add_paragraph("Offer Shares: 3,493,236,093")
    document.add_paragraph("Percentage Offered: 15%")
    document.add_paragraph("Nominal Value per Share: AED 1.00")
    document.add_paragraph("Offer Price Range: AED 1.30 – AED 1.50")
    document.add_paragraph("Alternative wording: AED 1.30 to AED 1.50")
    document.add_paragraph("Low/High values: AED 1.30 and AED 1.50")

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Issuer"
    table.rows[0].cells[1].text = "Talabat Holding plc"

    document.save(str(target))
    return target
